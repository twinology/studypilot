"""FastAPI backend for StudyPilot."""

import asyncio
import io
import logging
import logging.handlers
import os
import re
import shutil
import threading
from collections import deque
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from typing import Optional

import json

import hashlib
import urllib.parse
import urllib.request

import httpx
import uvicorn
from fastapi import FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

import config
from ai_provider import create_chat_completion, is_configured, load_settings, save_settings

# ── Logging setup ────────────────────────────────────────────────────
LOG_DIR = config.BASE_DIR / "logs"
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / "tutor.log"

# In-memory ring buffer for quick access via API
log_buffer: deque[dict] = deque(maxlen=500)


class BufferHandler(logging.Handler):
    """Custom handler that stores log records in an in-memory buffer."""
    def emit(self, record):
        log_buffer.append({
            "timestamp": datetime.fromtimestamp(record.created).strftime("%Y-%m-%d %H:%M:%S"),
            "level": record.levelname,
            "module": record.module,
            "message": record.getMessage(),
        })


# Configure tutor logger (propagate=False to avoid uvicorn duplication)
logger = logging.getLogger("tutor")
logger.setLevel(logging.INFO)
logger.propagate = False

# File handler (rotating, max 5MB, keep 3 backups)
file_handler = logging.handlers.RotatingFileHandler(
    LOG_FILE, maxBytes=5_000_000, backupCount=3, encoding="utf-8"
)
file_handler.setFormatter(logging.Formatter(
    "%(asctime)s | %(levelname)-7s | %(module)s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
))
logger.addHandler(file_handler)

# In-memory buffer handler
logger.addHandler(BufferHandler())

# Console handler
console_handler = logging.StreamHandler()
console_handler.setFormatter(logging.Formatter(
    "%(asctime)s | %(levelname)-7s | %(message)s",
    datefmt="%H:%M:%S",
))
logger.addHandler(console_handler)
from rag.chain import rag_query
from rag.chunker import chunk_text
from rag.document_loader import load_document
from rag.vector_store import (
    add_document,
    delete_document,
    list_documents,
    preload_model,
    reset_collection,
)
from tutor.feedback import generate_feedback
from tutor.scenarios import (
    generate_dynamic_scenario,
    get_scenario_by_id,
)
from tutor.session import SessionManager


# ── Conversation persistence ────────────────────────────────────────

def _save_conversation(conv_type: str, user_id: str, messages: list, metadata: dict = None):
    """Save a conversation (chat or session) to a JSON file."""
    import uuid as _uuid
    conv_id = datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + _uuid.uuid4().hex[:6]
    filename = f"{conv_type}_{conv_id}.json"
    data = {
        "id": conv_id,
        "type": conv_type,  # "chat" or "session"
        "user_id": user_id,
        "timestamp": datetime.now().isoformat(),
        "messages": messages,
        "metadata": metadata or {},
    }
    filepath = config.CONVERSATIONS_DIR / filename
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    logger.info(f"Gesprek opgeslagen: {filename} ({len(messages)} berichten)")
    return conv_id


def _load_conversations():
    """Load all saved conversations (metadata only)."""
    conversations = []
    for filepath in sorted(config.CONVERSATIONS_DIR.glob("*.json"), reverse=True):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            conversations.append({
                "id": data.get("id", filepath.stem),
                "type": data.get("type", "unknown"),
                "user_id": data.get("user_id", ""),
                "timestamp": data.get("timestamp", ""),
                "message_count": len(data.get("messages", [])),
                "metadata": data.get("metadata", {}),
                "filename": filepath.name,
            })
        except Exception as e:
            logger.error(f"Fout bij laden gesprek {filepath.name}: {e}")
    return conversations


def _load_conversation_detail(conv_id: str):
    """Load a single conversation with full messages."""
    for filepath in config.CONVERSATIONS_DIR.glob("*.json"):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            if data.get("id") == conv_id:
                return data
        except Exception:
            continue
    return None


def _start_telegram_bot():
    """Start Telegram bot in a separate thread."""
    if not config.TELEGRAM_BOT_TOKEN:
        logger.info("Geen TELEGRAM_BOT_TOKEN - Telegram bot niet gestart")
        return
    try:
        from bot.telegram_bot import run_bot
        thread = threading.Thread(target=run_bot, daemon=True)
        thread.start()
        logger.info("Telegram bot gestart")
    except Exception as e:
        logger.error(f"Telegram bot starten mislukt: {e}")


@asynccontextmanager
async def lifespan(app):
    port = os.environ.get("PORT", "8000")
    logger.info("=" * 50)
    logger.info("  StudyPilot")
    logger.info("=" * 50)
    logger.info(f"Web UI: http://localhost:{port}")
    logger.info(f"Documenten: {config.DOCUMENTS_DIR}")
    logger.info(f"ChromaDB: {config.CHROMA_DB_DIR}")
    settings = load_settings()
    logger.info(f"AI provider: {settings.get('provider', 'anthropic')}")
    logger.info(f"AI model: {settings.get('model', config.CLAUDE_MODEL)}")
    logger.info(f"Embedding model: {config.EMBEDDING_MODEL}")
    logger.info(f"Logbestand: {LOG_FILE}")
    _start_telegram_bot()
    yield


app = FastAPI(title="StudyPilot", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

session_manager = SessionManager()

# Store chat histories per user (simple in-memory)
chat_histories: dict[str, list[dict]] = {}

# RAG toggle: when False, chat and practice use Claude directly without knowledge base
rag_enabled: bool = True


# ── Request/Response models ──────────────────────────────────────────


class ChatRequest(BaseModel):
    message: str
    user_id: str = "web_user"


class SessionStartRequest(BaseModel):
    scenario_id: str
    user_id: str = "web_user"


class SessionMessageRequest(BaseModel):
    message: str
    user_id: str = "web_user"


class SessionFeedbackRequest(BaseModel):
    user_id: str = "web_user"


class TTSRequest(BaseModel):
    text: str
    voice: str = ""  # "male" of "female", leeg = default
    emotion: str = ""  # emotie voor stemkleur, bijv. "boos", "verdrietig"


class GenerateScenarioRequest(BaseModel):
    topic: Optional[str] = None
    difficulty: str = "basis"


# ── Document endpoints ───────────────────────────────────────────────


# Track background image processing progress
_image_progress: dict = {}  # filename -> {"total": N, "done": N, "status": str, "chunks": N}


def _save_doc_meta(filename: str, **kwargs):
    """Save metadata (index_mode, display_name) for a document in _crawl_meta.json."""
    meta_path = config.DOCUMENTS_DIR / "_crawl_meta.json"
    meta = {}
    if meta_path.exists():
        try:
            raw = json.loads(meta_path.read_text(encoding="utf-8"))
            # Migrate old format: {filename: url_string} → {filename: {display_name: url_string}}
            for k, v in raw.items():
                if isinstance(v, str):
                    meta[k] = {"display_name": v}
                elif isinstance(v, dict):
                    meta[k] = v
                else:
                    meta[k] = {}
        except Exception:
            pass
    if filename not in meta:
        meta[filename] = {}
    meta[filename].update(kwargs)
    meta_path.write_text(json.dumps(meta, indent=2, ensure_ascii=False), encoding="utf-8")


def _load_doc_meta() -> dict:
    """Load all document metadata from _crawl_meta.json."""
    meta_path = config.DOCUMENTS_DIR / "_crawl_meta.json"
    if not meta_path.exists():
        return {}
    try:
        raw = json.loads(meta_path.read_text(encoding="utf-8"))
        meta = {}
        for k, v in raw.items():
            if isinstance(v, str):
                meta[k] = {"display_name": v}
            elif isinstance(v, dict):
                meta[k] = v
            else:
                meta[k] = {}
        return meta
    except Exception:
        return {}


def _process_images_background(dest: Path, filename: str):
    """Process images in the background (runs in a thread). Called after text is already indexed."""
    try:
        from rag.image_extractor import extract_document_images
        from rag.image_describer import describe_images
        images = extract_document_images(dest, filename)
        if not images:
            _image_progress[filename] = {"total": 0, "done": 0, "status": "done", "chunks": 0}
            logger.info(f"Geen afbeeldingen gevonden in {filename}")
            return

        total = len(images)
        _image_progress[filename] = {"total": total, "done": 0, "status": "processing", "chunks": 0}

        # Process images one by one and update progress
        from rag.image_describer import describe_single_image
        all_chunks = []
        all_metas = []
        for i, img in enumerate(images):
            try:
                chunk, meta = describe_single_image(img, filename)
                if chunk:
                    all_chunks.append(chunk)
                    all_metas.append(meta)
            except Exception as e:
                logger.warning(f"Afbeelding {getattr(img, 'file_path', img)} overgeslagen: {e}")
            _image_progress[filename]["done"] = i + 1

        if all_chunks:
            num_image_chunks = add_document(all_chunks, filename, extra_metadatas=all_metas)
            _image_progress[filename]["chunks"] = num_image_chunks
            logger.info(f"✓ Afbeeldingen verwerkt: {filename} — {total} afbeeldingen, {num_image_chunks} chunks")
        else:
            logger.info(f"Geen afbeelding-chunks gegenereerd voor {filename}")

        _image_progress[filename]["status"] = "done"
    except Exception as e:
        _image_progress[filename] = {"total": 0, "done": 0, "status": "error", "chunks": 0, "error": str(e)}
        logger.warning(f"Image extractie fout op achtergrond: {e}")


@app.get("/api/upload/image-progress/{filename}")
async def get_image_progress(filename: str):
    """Get progress of background image processing."""
    if filename in _image_progress:
        result = _image_progress[filename]
    else:
        result = {"total": 0, "done": 0, "status": "unknown", "chunks": 0}
    return JSONResponse(content=result, headers={
        "Cache-Control": "no-cache, no-store, must-revalidate",
        "Pragma": "no-cache",
        "Expires": "0",
    })


@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...), index_mode: str = Form("all")):
    """Upload and process a document. index_mode: 'all' (text+images), 'text' (text only), 'images' (images only)."""
    allowed = {".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".htm"}
    suffix = Path(file.filename).suffix.lower()
    if suffix not in allowed:
        raise HTTPException(400, f"Niet-ondersteund formaat: {suffix}. Toegestaan: {', '.join(allowed)}")

    # Save file
    dest = config.DOCUMENTS_DIR / file.filename
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    do_text = index_mode in ("all", "text")
    do_images = index_mode in ("all", "images")

    # Process text
    num_chunks = 0
    text = ""
    try:
        if do_text:
            text = load_document(dest)
            if not text or not text.strip():
                if not do_images:
                    os.remove(dest)
                    raise HTTPException(400, "Kon geen tekst uit het document extraheren.")
            else:
                chunks = chunk_text(text)
                num_chunks = add_document(chunks, file.filename)
                logger.info(f"Document geupload: {file.filename} ({len(text)} tekens, {num_chunks} chunks, modus={index_mode})")
        else:
            logger.info(f"Document geupload (alleen afbeeldingen): {file.filename} (modus={index_mode})")

        # Schedule image processing in background thread (won't block the response)
        images_scheduled = False
        if do_images and config.EXTRACT_IMAGES and suffix in (".pdf", ".docx", ".doc"):
            thread = threading.Thread(target=_process_images_background, args=(dest, file.filename), daemon=True)
            thread.start()
            images_scheduled = True
            logger.info(f"Afbeelding-verwerking gestart op achtergrond voor {file.filename}")

        # Save index_mode in meta
        _save_doc_meta(file.filename, index_mode=index_mode)

        return {
            "status": "ok",
            "filename": file.filename,
            "chunks": num_chunks,
            "image_chunks": 0,
            "images_processing": images_scheduled,
            "index_mode": index_mode,
            "characters": len(text),
        }
    except ValueError as e:
        if dest.exists():
            os.remove(dest)
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.error(f"Onverwachte fout bij verwerking van {file.filename}: {e}")
        raise HTTPException(500, f"Fout bij verwerking: {str(e)}")


class UrlUploadRequest(BaseModel):
    url: str
    index_mode: str = "all"  # "all", "text", "images"


_FETCH_UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
_FETCH_HEADERS = {
    "User-Agent": _FETCH_UA,
    "Accept": "text/html,application/xhtml+xml,*/*;q=0.8",
    "Accept-Language": "nl,en;q=0.5",
}


def _fetch_url_sync(url: str) -> tuple:
    """Fetch URL using urllib (fallback for sites that block httpx). Returns (html_content, content_type)."""
    req = urllib.request.Request(url, headers=_FETCH_HEADERS)
    with urllib.request.urlopen(req, timeout=30) as resp:
        content_type = resp.headers.get("Content-Type", "")
        raw = resp.read()
        encoding = "utf-8"
        if "charset=" in content_type:
            encoding = content_type.split("charset=")[-1].split(";")[0].strip()
        return raw.decode(encoding, errors="replace"), content_type


@app.post("/api/upload-url")
async def upload_url(req: UrlUploadRequest):
    """Fetch a webpage, extract text content, save as HTML and index it."""
    from urllib.parse import urlparse as _urlparse

    url = req.url.strip()
    if not url:
        raise HTTPException(400, "URL mag niet leeg zijn.")

    # Validate URL
    parsed = _urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise HTTPException(400, "Alleen HTTP en HTTPS URLs zijn toegestaan.")

    # Fetch the webpage — try httpx first, fall back to urllib (some sites block httpx TLS fingerprint)
    html_content = None
    content_type = ""

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
            response = await client.get(url, headers=_FETCH_HEADERS)
            response.raise_for_status()
            content_type = response.headers.get("content-type", "")
            html_content = response.text
    except Exception:
        # Fallback to urllib (works with sites that block httpx, like Wikipedia)
        try:
            html_content, content_type = await asyncio.to_thread(_fetch_url_sync, url)
        except Exception as e:
            raise HTTPException(502, f"Kon de URL niet ophalen: {str(e)}")

    # Check content type
    if "text/html" not in content_type and "text/plain" not in content_type:
        raise HTTPException(400, f"Niet-ondersteund content type: {content_type}. Alleen HTML en tekst pagina's worden ondersteund.")

    if not html_content or not html_content.strip():
        raise HTTPException(400, "De pagina bevat geen inhoud.")

    # Generate a filename from the URL
    # Use domain + path, sanitized for filesystem
    domain = parsed.netloc.replace("www.", "")
    path_part = parsed.path.strip("/").replace("/", "_")
    if not path_part:
        path_part = "index"
    # Create a short hash to avoid collisions
    url_hash = hashlib.md5(url.encode()).hexdigest()[:6]
    safe_name = re.sub(r'[^\w\-.]', '_', f"{domain}_{path_part}")
    # Truncate if too long
    if len(safe_name) > 80:
        safe_name = safe_name[:80]
    filename = f"{safe_name}_{url_hash}.html"

    # Save the HTML file
    dest = config.DOCUMENTS_DIR / filename
    with open(dest, "w", encoding="utf-8") as f:
        f.write(html_content)

    # Process and index
    try:
        text = load_document(dest)
        if not text or not text.strip():
            os.remove(dest)
            raise HTTPException(400, "Kon geen tekst uit de webpagina extraheren.")
        chunks = chunk_text(text)
        num_chunks = add_document(chunks, filename)
        logger.info(f"URL geïndexeerd: {url} -> {filename} ({len(text)} tekens, {num_chunks} chunks)")

        # Save metadata
        _save_doc_meta(filename, display_name=url, index_mode=req.index_mode if hasattr(req, 'index_mode') else "text")

        return {
            "status": "ok",
            "filename": filename,
            "url": url,
            "chunks": num_chunks,
            "characters": len(text),
        }
    except ValueError as e:
        if dest.exists():
            os.remove(dest)
        raise HTTPException(400, str(e))


class CrawlWebsiteRequest(BaseModel):
    url: str
    max_depth: int = 3
    max_pages: int = 50
    index_mode: str = "all"  # "all", "text", "images"


@app.post("/api/crawl-website")
async def crawl_website(req: CrawlWebsiteRequest):
    """Crawl an entire website using Tavily and index all pages into the knowledge base."""
    from urllib.parse import urlparse as _urlparse

    url = req.url.strip()
    if not url:
        raise HTTPException(400, "URL mag niet leeg zijn.")

    parsed = _urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise HTTPException(400, "Alleen HTTP en HTTPS URLs zijn toegestaan.")

    # Get Tavily API key from settings
    settings = load_settings()
    tavily_key = settings.get("tavily_api_key", "").strip()
    if not tavily_key:
        raise HTTPException(400, "Tavily API key is niet geconfigureerd. Stel deze in bij Setup.")

    try:
        from tavily import TavilyClient
    except ImportError:
        raise HTTPException(500, "tavily-python is niet geïnstalleerd. Run: pip install tavily-python")

    domain = parsed.netloc.replace("www.", "")

    # Determine path filter: only crawl pages at and under the given URL path
    base_path = parsed.path.rstrip("/")
    if base_path and base_path != "/":
        select_paths = [base_path, base_path + "/*"]
    else:
        select_paths = None

    # Run Tavily crawl in a thread to not block the event loop
    def _do_crawl():
        client = TavilyClient(api_key=tavily_key)
        crawl_kwargs = dict(
            url=url,
            max_depth=req.max_depth,
            max_breadth=20,
            limit=req.max_pages,
            format="markdown",
            extract_depth="advanced",
        )
        if select_paths:
            crawl_kwargs["select_paths"] = select_paths
        result = client.crawl(**crawl_kwargs)
        return result

    try:
        crawl_result = await asyncio.to_thread(_do_crawl)
    except Exception as e:
        error_msg = str(e)
        if "401" in error_msg or "unauthorized" in error_msg.lower():
            raise HTTPException(401, "Tavily API key is ongeldig of verlopen.")
        raise HTTPException(502, f"Tavily crawl fout: {error_msg}")

    # Process crawled pages — combine all into a single document
    results = crawl_result.get("results", [])
    if not results:
        raise HTTPException(400, "Geen pagina's gevonden op deze website.")

    def _fetch_full_page(page_url: str) -> str:
        """Fetch a page ourselves and extract body text when Tavily truncates content."""
        import urllib.request
        from html.parser import HTMLParser

        class _BodyTextExtractor(HTMLParser):
            """Extract meaningful text from HTML, skipping nav/header/footer/script/style."""
            SKIP_TAGS = {'script', 'style', 'nav', 'header', 'footer', 'noscript', 'iframe'}

            def __init__(self):
                super().__init__()
                self._text_parts = []
                self._skip_depth = 0
                self._tag_stack = []

            def handle_starttag(self, tag, attrs):
                attrs_dict = dict(attrs)
                self._tag_stack.append(tag)
                # Skip nav/header/footer/script/style elements
                if tag in self.SKIP_TAGS:
                    self._skip_depth += 1
                    return
                # Skip elements with nav/menu/footer classes
                cls = attrs_dict.get('class', '').lower()
                if any(k in cls for k in ['nav', 'menu', 'footer', 'header', 'sidebar', 'cookie']):
                    self._skip_depth += 1
                    return
                # Add line breaks for block elements
                if tag in ('p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'br', 'tr', 'blockquote'):
                    self._text_parts.append('\n')
                # Add markdown-style headings
                if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                    level = int(tag[1])
                    self._text_parts.append('#' * level + ' ')

            def handle_endtag(self, tag):
                if self._tag_stack and self._tag_stack[-1] == tag:
                    self._tag_stack.pop()
                if tag in self.SKIP_TAGS:
                    self._skip_depth = max(0, self._skip_depth - 1)
                # Check for class-based skip (approximate)
                if self._skip_depth > 0 and tag in ('div', 'section', 'aside'):
                    self._skip_depth = max(0, self._skip_depth - 1)

            def handle_data(self, data):
                if self._skip_depth > 0:
                    return
                text = data.strip()
                if text:
                    self._text_parts.append(text)

            def get_text(self):
                return '\n'.join(self._text_parts)

        try:
            req = urllib.request.Request(page_url, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'
            })
            resp = urllib.request.urlopen(req, timeout=15)
            html = resp.read().decode('utf-8', errors='replace')
            parser = _BodyTextExtractor()
            parser.feed(html)
            text = parser.get_text()
            # Clean up excessive whitespace
            text = re.sub(r'\n{3,}', '\n\n', text)
            return text.strip()
        except Exception as e:
            logger.warning(f"Fallback fetch mislukt voor {page_url}: {e}")
            return ""

    def _clean_crawled_content(text: str) -> str:
        """Strip navigation menus, headers and footers — keep only the body content."""
        lines = text.split("\n")
        cleaned = []
        body_started = False
        nav_link_count = 0  # Count consecutive nav links

        # --- Phase 1: find where the real body content starts ---
        # The header consists of consecutive nav links, boilerplate, phone numbers, etc.
        # Once we hit a heading (#) or substantial paragraph text, the body has started.
        header_end = 0
        consecutive_nav = 0
        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue
            is_nav_link = bool(re.match(r'^[\*\-\+]\s+\[.+?\]\(', stripped))
            is_sub_link = bool(re.match(r'^\+\s+\[.+?\]\(', stripped))
            is_boilerplate = bool(re.match(
                r'^\[(Bel (gratis )?met:|Opleiderscore:)', stripped
            )) or bool(re.match(
                r'^\* \[(Alle trainingen|Goed bereikbare|Maatwerk in leer|Ervaren trainers|Eindevaluatie|Effectieve|Praktijkgerichte|Veilige)', stripped
            ))
            is_breadcrumb = stripped.startswith("* [Home]")

            if is_nav_link or is_sub_link or is_boilerplate or is_breadcrumb:
                consecutive_nav += 1
                header_end = i + 1
            elif consecutive_nav >= 3:
                # We've seen 3+ consecutive nav links followed by non-nav content
                # This is where the body starts
                header_end = i
                break
            else:
                # Non-nav content before seeing enough nav links — likely body already
                break

        # --- Phase 2: find where the footer starts ---
        footer_start = len(lines)
        for i in range(len(lines) - 1, -1, -1):
            stripped = lines[i].strip()
            if not stripped:
                continue
            if re.match(r'^[©�]\s*\d{4}', stripped) or re.match(r'^Wij verlenen onze diensten', stripped):
                footer_start = i
                continue
            if footer_start < len(lines) and i < footer_start:
                break

        # --- Phase 3: extract body content between header and footer ---
        for i in range(header_end, footer_start):
            line = lines[i]
            stripped = line.strip()

            # Skip breadcrumb lines that may appear right after the header
            if not body_started:
                if not stripped:
                    continue
                if stripped.startswith("* [Home]"):
                    continue
                if re.match(r'^\[(Bel (gratis )?met:|Opleiderscore:)', stripped):
                    continue

            body_started = True
            cleaned.append(line)

        # Strip trailing empty lines
        while cleaned and not cleaned[-1].strip():
            cleaned.pop()

        return "\n".join(cleaned)

    # Collect all page contents into one combined markdown document
    combined_parts = [f"# Website crawl: {url}\n"]
    page_details = []
    errors = []

    for page in results:
        page_url = page.get("url", "")
        raw_content = page.get("raw_content", "") or page.get("content", "")
        if not raw_content or not raw_content.strip():
            continue

        # Clean content: remove nav, header, footer
        cleaned_content = _clean_crawled_content(raw_content)

        # If Tavily truncated the content, fetch ourselves for full text
        # (Tavily often returns ~15KB while pages have 40KB+ of content)
        if len(raw_content) < 20000 and page_url:
            fallback_text = _fetch_full_page(page_url)
            if fallback_text and len(fallback_text) > len(cleaned_content) * 1.2:
                logger.info(f"Fallback fetch voor {page_url}: Tavily {len(cleaned_content)} -> eigen {len(fallback_text)} tekens")
                cleaned_content = fallback_text

        if not cleaned_content.strip():
            continue

        # Get page title: prefer largest heading (# or ##), skip small ###### headings
        page_title = ""
        for line in cleaned_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            # Prefer # or ## headings (main titles)
            if re.match(r'^#{1,2}\s+', stripped):
                candidate = stripped.lstrip("#").strip()
                if len(candidate) > 5:  # Skip very short/incomplete titles
                    page_title = candidate
                    break
        if not page_title:
            # Fallback: use URL path as readable title
            page_parsed = _urlparse(page_url)
            path_parts = page_parsed.path.strip("/").split("/")
            page_title = " — ".join(p.replace("-", " ").title() for p in path_parts if p) or domain

        combined_parts.append(f"\n\n---\n\n## Bron: {page_url}\n\n{cleaned_content}")
        page_details.append({
            "url": page_url,
            "title": page_title,
            "characters": len(cleaned_content),
            "content": cleaned_content,
        })

    if not page_details:
        raise HTTPException(400, "Kon geen content extraheren van deze website.")

    combined_text = "\n".join(combined_parts)

    # Generate a single filename for the entire crawl
    path_part = parsed.path.strip("/").replace("/", "_")
    if not path_part:
        path_part = "site"
    url_hash = hashlib.md5(url.encode()).hexdigest()[:6]
    safe_name = re.sub(r'[^\w\-.]', '_', f"{domain}_{path_part}")
    if len(safe_name) > 80:
        safe_name = safe_name[:80]
    filename = f"{safe_name}_{url_hash}.md"

    # Save the combined document
    dest = config.DOCUMENTS_DIR / filename
    try:
        with open(dest, "w", encoding="utf-8") as f:
            f.write(combined_text)

        # Save metadata mapping filename -> original URL for display
        _save_doc_meta(filename, display_name=url, index_mode=index_mode)

        text = load_document(dest)
        if not text or not text.strip():
            os.remove(dest)
            raise HTTPException(400, "Kon geen tekst uit de gecrawlde pagina's extraheren.")

        chunks = chunk_text(text)
        total_chunks = add_document(chunks, filename)
        total_chars = len(text)

        logger.info(f"Website gecrawld: {url} -> {filename} ({len(page_details)} pagina's, {total_chunks} chunks, {total_chars} tekens)")

    except ValueError as e:
        if dest.exists():
            os.remove(dest)
        raise HTTPException(400, str(e))

    return {
        "status": "ok",
        "base_url": url,
        "filename": filename,
        "pages_found": len(results),
        "pages_indexed": len(page_details),
        "total_chunks": total_chunks,
        "total_characters": total_chars,
        "pages": page_details,
        "errors": errors,
    }


@app.get("/api/documents")
async def get_documents():
    """List all uploaded documents (files on disk + index status)."""
    allowed = {".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".htm"}

    # Get indexed documents from vector store (may be slow on first call)
    try:
        indexed_docs = set(list_documents())
    except Exception:
        indexed_docs = set()

    # Load document metadata (display names, index modes)
    doc_meta = _load_doc_meta()

    # List all physical files in the documents directory
    file_info = []
    if config.DOCUMENTS_DIR.exists():
        for f in sorted(config.DOCUMENTS_DIR.iterdir()):
            if f.is_file() and f.suffix.lower() in allowed:
                meta = doc_meta.get(f.name, {})
                doc = {
                    "name": f.name,
                    "size": f.stat().st_size,
                    "indexed": f.name in indexed_docs,
                    "index_mode": meta.get("index_mode", "all"),
                }
                if "display_name" in meta:
                    doc["display_name"] = meta["display_name"]
                file_info.append(doc)
    # Include any active image processing status
    active_processing = {}
    for fname, prog in _image_progress.items():
        if prog.get("status") == "processing":
            active_processing[fname] = prog

    return {"documents": file_info, "image_processing": active_processing}


@app.delete("/api/documents/{doc_name}")
async def remove_document(doc_name: str):
    """Remove a document from the system."""
    deleted = delete_document(doc_name)
    file_path = config.DOCUMENTS_DIR / doc_name
    if file_path.exists():
        os.remove(file_path)
    # Clean up extracted images
    image_dir = config.IMAGES_DIR / Path(doc_name).stem
    if image_dir.exists():
        shutil.rmtree(image_dir)
    # Clean up crawl metadata
    meta_path = config.DOCUMENTS_DIR / "_crawl_meta.json"
    if meta_path.exists():
        try:
            crawl_meta = json.loads(meta_path.read_text(encoding="utf-8"))
            if doc_name in crawl_meta:
                del crawl_meta[doc_name]
                meta_path.write_text(json.dumps(crawl_meta, indent=2, ensure_ascii=False), encoding="utf-8")
        except Exception:
            pass
    return {"status": "ok", "chunks_deleted": deleted}


class ReindexRequest(BaseModel):
    index_mode: str = "all"


@app.post("/api/documents/{doc_name}/reindex")
async def reindex_document(doc_name: str, req: ReindexRequest):
    """Re-index a single document with the specified index mode."""
    file_path = config.DOCUMENTS_DIR / doc_name
    if not file_path.exists():
        raise HTTPException(404, "Document niet gevonden.")

    suffix = file_path.suffix.lower()
    do_text = req.index_mode in ("all", "text")
    do_images = req.index_mode in ("all", "images")

    # Remove old index entries first
    try:
        delete_document(doc_name)
    except Exception:
        pass

    # Re-index text
    num_chunks = 0
    if do_text:
        text = load_document(file_path)
        if text and text.strip():
            chunks = chunk_text(text)
            num_chunks = add_document(chunks, doc_name)
            logger.info(f"Document hergeïndexeerd: {doc_name} ({len(text)} tekens, {num_chunks} chunks, modus={req.index_mode})")

    # Schedule image re-processing
    images_scheduled = False
    if do_images and config.EXTRACT_IMAGES and suffix in (".pdf", ".docx", ".doc"):
        # Clean old images first
        image_dir = config.IMAGES_DIR / file_path.stem
        if image_dir.exists():
            shutil.rmtree(image_dir)
        thread = threading.Thread(target=_process_images_background, args=(file_path, doc_name), daemon=True)
        thread.start()
        images_scheduled = True
        logger.info(f"Afbeelding-herverwerking gestart op achtergrond voor {doc_name}")

    # Update index mode in metadata
    _save_doc_meta(doc_name, index_mode=req.index_mode)

    return {
        "status": "ok",
        "filename": doc_name,
        "chunks": num_chunks,
        "images_processing": images_scheduled,
        "index_mode": req.index_mode,
    }


@app.get("/api/documents/{doc_name}/content")
async def get_document_content(doc_name: str):
    """Return the text content of a document for viewing."""
    file_path = config.DOCUMENTS_DIR / doc_name
    if not file_path.exists():
        raise HTTPException(404, "Document niet gevonden.")
    try:
        text = file_path.read_text(encoding="utf-8")
    except Exception:
        try:
            text = file_path.read_text(encoding="latin-1")
        except Exception as e:
            raise HTTPException(500, f"Kan document niet lezen: {e}")

    # Load display name from document metadata
    doc_meta = _load_doc_meta()
    meta = doc_meta.get(doc_name, {})
    display_name = meta.get("display_name", doc_name)

    return {"name": doc_name, "display_name": display_name, "content": text, "size": len(text)}


# ── Re-index endpoint ─────────────────────────────────────────────────


def _reindex_sync():
    """Synchronous reindex (runs in thread to avoid blocking event loop)."""
    doc_dir = config.DOCUMENTS_DIR
    files = [f for f in doc_dir.iterdir() if f.is_file() and f.suffix.lower() in {
        ".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".htm"
    }]

    if not files:
        return {"status": "ok", "message": "Geen documenten gevonden om te herindexeren.", "documents": 0}

    logger.info(f"Herindexering gestart: {len(files)} documenten")

    # Reset the vector store (delete old embeddings)
    reset_collection()

    # Clean up old images
    if config.IMAGES_DIR.exists():
        shutil.rmtree(config.IMAGES_DIR)
        config.IMAGES_DIR.mkdir(exist_ok=True)

    results = []
    total_chunks = 0
    for file_path in files:
        try:
            text = load_document(file_path)
            if text and text.strip():
                chunks = chunk_text(text)
                num_chunks = add_document(chunks, file_path.name)
                total_chunks += num_chunks
                results.append({"name": file_path.name, "chunks": num_chunks, "status": "ok"})
                logger.info(f"Herindexering: {file_path.name} → {num_chunks} chunks")

                # Multimodal: extract and describe images
                if config.EXTRACT_IMAGES and file_path.suffix.lower() in (".pdf", ".docx", ".doc"):
                    try:
                        from rag.image_extractor import extract_document_images
                        from rag.image_describer import describe_images
                        images = extract_document_images(file_path, file_path.name)
                        if images:
                            img_chunks, img_metas = describe_images(images, file_path.name)
                            if img_chunks:
                                add_document(img_chunks, file_path.name, extra_metadatas=img_metas)
                                total_chunks += len(img_chunks)
                                logger.info(f"Herindexering afbeeldingen: {file_path.name} → {len(img_chunks)} image chunks")
                    except Exception as e:
                        logger.warning(f"Image extractie fout bij herindexering ({file_path.name}): {e}")
            else:
                results.append({"name": file_path.name, "chunks": 0, "status": "leeg"})
        except Exception as e:
            results.append({"name": file_path.name, "chunks": 0, "status": f"fout: {e}"})
            logger.error(f"Herindexering fout: {file_path.name}: {e}")

    logger.info(f"Herindexering voltooid: {len(files)} documenten, {total_chunks} chunks totaal")
    return {
        "status": "ok",
        "documents": len(files),
        "total_chunks": total_chunks,
        "details": results,
    }


@app.post("/api/reindex")
async def reindex_documents():
    """Re-index all documents with current embedding model. Use after model change."""
    result = await asyncio.to_thread(_reindex_sync)
    return result


# ── RAG toggle endpoint ──────────────────────────────────────────────


@app.get("/api/rag-toggle")
async def get_rag_toggle():
    """Get the current RAG toggle state."""
    return {"enabled": rag_enabled}


@app.post("/api/rag-toggle")
async def set_rag_toggle():
    """Toggle RAG on/off. When off, chat and practice use Claude directly."""
    global rag_enabled
    rag_enabled = not rag_enabled
    state = "ingeschakeld" if rag_enabled else "uitgeschakeld"
    logger.info(f"Kennisbank {state}")
    return {"enabled": rag_enabled}


# ── Chat endpoint ────────────────────────────────────────────────────


@app.post("/api/chat")
async def chat(req: ChatRequest):
    """Regular RAG chat about conversation techniques."""
    if not is_configured():
        raise HTTPException(503, "Geen AI API-key geconfigureerd. Ga naar Setup.")
    # Check if user has an active session
    if session_manager.has_active_session(req.user_id):
        raise HTTPException(400, "Je hebt een actieve oefensessie. Beëindig deze eerst met /api/session/feedback.")

    history = chat_histories.get(req.user_id, [])
    mode = "RAG" if rag_enabled else "Direct"
    logger.info(f"Chat [{req.user_id}] ({mode}): {req.message[:80]}...")
    try:
        result = rag_query(req.message, chat_history=history[-10:], use_rag=rag_enabled)
        usage = result.get("usage", {})
        logger.info(
            f"Chat antwoord [{req.user_id}] ({mode}): {result['answer'][:80]}... "
            f"| tokens in={usage.get('input_tokens', '?')} out={usage.get('output_tokens', '?')} "
            f"chunks={usage.get('context_chunks', '?')}"
        )
    except RuntimeError as e:
        logger.error(f"Chat fout [{req.user_id}]: {e}")
        raise HTTPException(503, str(e))

    # Update history
    if req.user_id not in chat_histories:
        chat_histories[req.user_id] = []
    chat_histories[req.user_id].append({"role": "user", "content": req.message})
    chat_histories[req.user_id].append({"role": "assistant", "content": result["answer"]})

    return result


# ── Scenario endpoints ───────────────────────────────────────────────


@app.post("/api/scenarios/generate")
async def create_dynamic_scenario(req: GenerateScenarioRequest):
    """Generate a new scenario from document content."""
    if not is_configured():
        raise HTTPException(503, "Geen AI API-key geconfigureerd. Ga naar Setup.")
    logger.info(f"Scenario genereren: thema='{req.topic or 'geen'}', niveau={req.difficulty}, rag={rag_enabled}")
    scenario = await generate_dynamic_scenario(req.topic, req.difficulty, use_rag=rag_enabled)
    logger.info(f"Scenario gegenereerd: {scenario.name} ({scenario.difficulty})")
    return {"scenario": scenario.to_dict()}


# ── Session endpoints ────────────────────────────────────────────────


@app.post("/api/session/start")
async def start_session(req: SessionStartRequest):
    """Start a practice session with a scenario."""
    if session_manager.has_active_session(req.user_id):
        # Automatically end the previous session so the student can start fresh
        session_manager.end_session(req.user_id)
        logger.info(f"Vorige sessie automatisch beëindigd [{req.user_id}]")

    scenario = get_scenario_by_id(req.scenario_id)
    if not scenario:
        raise HTTPException(404, f"Scenario '{req.scenario_id}' niet gevonden.")

    session = session_manager.start_session(req.user_id, scenario)
    logger.info(f"Sessie gestart [{req.user_id}]: {scenario.name} ({scenario.difficulty})")

    # Generate opening message from tutor in role
    if not is_configured():
        raise HTTPException(503, "Geen AI API-key geconfigureerd. Ga naar Setup.")
    char_name = scenario.character_name or "het personage"
    system = f"""Je bent een acteur die een rol speelt voor een oefengesprek.
Je personage heet: {char_name}

## Je rol:
{scenario.role_instruction}

## Belangrijk:
- Blijf ALTIJD in je rol
- Reageer realistisch en consistent met je karakter
- Communiceer in het Nederlands
- Begin het gesprek vanuit je rol (stel je voor of begin de situatie)
- Markeer ELKE actie, gebaar of toonverandering met het [CONTEXT] prefix op een eigen regel.
  Formaat: [CONTEXT] {char_name} <beschrijving>
  Gebruik dit niet alleen aan het begin, maar ook TUSSENDOOR wanneer het personage iets doet.
  Gebruik NOOIT cursieve tekst (*...*) voor acties — altijd [CONTEXT].
  Na elke [CONTEXT] regel volgt een lege regel en dan de gesproken tekst."""

    # Add user profile context
    user_name = load_settings().get("user_name", "").strip()
    if user_name:
        system += f"\n- De student heet {user_name}. Spreek de student aan met '{user_name}' en gebruik 'je' en 'jij'."

    try:
        opening = create_chat_completion(
            messages=[{"role": "user", "content": "Start het gesprek vanuit je rol."}],
            system=system,
            max_tokens=config.MAX_TOKENS,
        )
    except RuntimeError as e:
        session_manager.end_session(req.user_id)
        raise HTTPException(503, str(e))
    except Exception as e:
        session_manager.end_session(req.user_id)
        raise HTTPException(503, f"AI-fout: {e}")
    session.add_message("tutor", opening)

    return {
        "session_id": session.session_id,
        "scenario": scenario.to_dict(),
        "opening_message": opening,
    }


@app.post("/api/session/message")
async def session_message(req: SessionMessageRequest):
    """Send a message in an active practice session."""
    session = session_manager.get_session(req.user_id)
    if not session:
        raise HTTPException(400, "Geen actieve oefensessie. Start er eerst een.")

    session.add_message("student", req.message)

    # Get RAG context for better in-character responses (only when RAG is enabled)
    context_hint = ""
    if rag_enabled:
        from rag.vector_store import search as rag_search
        context_items = rag_search(req.message, top_k=3)
        if context_items:
            context_hint = "\n\n[Interne hint - niet benoemen in je antwoord: " + " ".join(
                item["text"][:200] for item in context_items
            ) + "]"

    char_name = session.scenario.character_name or "het personage"
    system = f"""Je bent een acteur die een rol speelt voor een oefengesprek.
Je personage heet: {char_name}

## Je rol:
{session.scenario.role_instruction}

## Belangrijk:
- Blijf ALTIJD in je rol
- Reageer realistisch en consistent met je karakter
- Communiceer in het Nederlands
- Reageer op wat de student zegt vanuit je rol
- Houd je antwoorden beknopt (2-4 zinnen), zoals in een echt gesprek
- Markeer ELKE actie, gebaar of toonverandering met het [CONTEXT] prefix op een eigen regel.
  Formaat: [CONTEXT] {char_name} <beschrijving>
  Gebruik dit niet alleen aan het begin, maar ook TUSSENDOOR wanneer het personage iets doet.
  Gebruik NOOIT cursieve tekst (*...*) voor acties — altijd [CONTEXT].
  Na elke [CONTEXT] regel volgt een lege regel en dan de gesproken tekst.{context_hint}"""

    # Add user profile context
    user_name = load_settings().get("user_name", "").strip()
    if user_name:
        system += f"\n- De student heet {user_name}. Spreek de student aan met '{user_name}' en gebruik 'je' en 'jij'."

    try:
        result = create_chat_completion(
            messages=session.to_claude_messages(),
            system=system,
            max_tokens=512,
            return_usage=True,
        )
        tutor_reply = result["text"]
        token_usage = result["usage"]
    except RuntimeError as e:
        session.messages.pop()
        raise HTTPException(503, str(e))
    except Exception as e:
        session.messages.pop()
        raise HTTPException(503, f"AI-fout: {e}")
    session.add_message("tutor", tutor_reply)
    logger.info(f"Sessie bericht [{req.user_id}]: student={req.message[:50]}... | tutor={tutor_reply[:50]}...")

    return {"reply": tutor_reply, "message_count": len(session.messages), "usage": token_usage}


@app.post("/api/session/feedback")
async def session_feedback(req: SessionFeedbackRequest):
    """End session and generate detailed feedback."""
    session = session_manager.end_session(req.user_id)
    if not session:
        raise HTTPException(400, "Geen actieve oefensessie om te beëindigen.")

    if len(session.messages) < 2:
        return {
            "feedback": "Er zijn te weinig berichten uitgewisseld voor een zinvolle analyse. Probeer de gesprekssimulatie langer voort te zetten.",
            "message_count": len(session.messages),
        }

    feedback = await generate_feedback(session, use_rag=rag_enabled)
    session.feedback = feedback
    logger.info(f"Feedback gegenereerd [{req.user_id}]: {len(session.messages)} berichten")

    # Save completed session with feedback
    session_messages = [
        {"role": m.role, "content": m.content} for m in session.messages
    ]
    _save_conversation(
        "session", req.user_id,
        session_messages,
        {
            "scenario_id": session.scenario.id,
            "scenario_name": session.scenario.name,
            "scenario_difficulty": session.scenario.difficulty,
            "feedback": feedback,
        },
    )

    return {
        "feedback": feedback,
        "message_count": len(session.messages),
        "scenario": session.scenario.to_dict(),
    }


# ── TTS endpoint ─────────────────────────────────────────────────────


@app.post("/api/tts")
async def text_to_speech(req: TTSRequest):
    """Convert text to speech using ElevenLabs."""
    el_key = load_settings().get("elevenlabs_api_key", "") or config.ELEVENLABS_API_KEY
    if not el_key:
        raise HTTPException(400, "Geen ElevenLabs API key. Browser TTS wordt gebruikt.")

    from elevenlabs import ElevenLabs, VoiceSettings

    try:
        # Select voice based on gender
        voice_key = req.voice if req.voice in config.ELEVENLABS_VOICES else config.ELEVENLABS_DEFAULT_VOICE
        voice_id = config.ELEVENLABS_VOICES[voice_key]

        # Apply emotion-based voice settings
        emotion_cfg = config.ELEVENLABS_EMOTION_SETTINGS.get(
            req.emotion, config.ELEVENLABS_EMOTION_SETTINGS["neutraal"]
        )
        voice_settings = VoiceSettings(
            stability=emotion_cfg["stability"],
            similarity_boost=emotion_cfg["similarity_boost"],
            style=emotion_cfg["style"],
            use_speaker_boost=emotion_cfg["use_speaker_boost"],
        )

        client = ElevenLabs(api_key=el_key)
        audio_generator = client.text_to_speech.convert(
            voice_id=voice_id,
            model_id=config.ELEVENLABS_MODEL_ID,
            text=req.text,
            voice_settings=voice_settings,
        )

        # Collect audio bytes from generator
        audio_bytes = b"".join(audio_generator)

        return StreamingResponse(
            iter([audio_bytes]),
            media_type="audio/mpeg",
            headers={"Content-Disposition": "inline"},
        )
    except Exception as e:
        error_msg = str(e)
        logger.error(f"TTS fout: {error_msg}")
        if "text_to_speech" in error_msg.lower() and "permission" in error_msg.lower():
            raise HTTPException(
                403,
                "Je ElevenLabs API-key mist de 'text_to_speech' permissie. "
                "Ga naar https://elevenlabs.io/app/settings/api-keys en maak "
                "een nieuwe key aan met de juiste permissies.",
            )
        raise HTTPException(500, f"TTS fout: {error_msg}")


# ── Logs endpoint ────────────────────────────────────────────────────


@app.get("/api/logs")
async def get_logs(
    limit: int = Query(100, ge=1, le=500),
    level: Optional[str] = Query(None),
):
    """Return recent log entries from the in-memory buffer."""
    entries = list(log_buffer)
    if level:
        level_upper = level.upper()
        entries = [e for e in entries if e["level"] == level_upper]
    return {"logs": entries[-limit:], "total": len(entries)}


@app.get("/api/logs/file")
async def get_log_file():
    """Download the full log file."""
    if not LOG_FILE.exists():
        raise HTTPException(404, "Logbestand niet gevonden.")
    return FileResponse(LOG_FILE, filename="tutor.log", media_type="text/plain")


# ── Conversation history endpoints ──────────────────────────────────


@app.get("/api/conversations")
async def get_conversations(conv_type: Optional[str] = Query(None)):
    """List all saved conversations."""
    conversations = _load_conversations()
    if conv_type:
        conversations = [c for c in conversations if c["type"] == conv_type]
    return {"conversations": conversations}


@app.get("/api/conversations/{conv_id}")
async def get_conversation_detail(conv_id: str):
    """Get full conversation detail."""
    data = _load_conversation_detail(conv_id)
    if not data:
        raise HTTPException(404, "Rapport niet gevonden.")
    return data


@app.get("/api/conversations/{conv_id}/export")
async def export_conversation(conv_id: str, format: str = Query("md")):
    """Export a conversation as MD, DOCX, or PDF."""
    if format not in ("md", "docx", "pdf"):
        raise HTTPException(400, "Ongeldig formaat. Kies 'md', 'docx' of 'pdf'.")

    data = _load_conversation_detail(conv_id)
    if not data:
        raise HTTPException(404, "Rapport niet gevonden.")

    conv_type = data.get("type", "rapport")
    meta = data.get("metadata", {})
    title = meta.get("scenario_name") or meta.get("topic") or "Rapport"
    timestamp = data.get("timestamp", "")
    messages = data.get("messages", [])
    difficulty = meta.get("scenario_difficulty", "") or meta.get("quiz_difficulty", "")
    feedback = meta.get("feedback", "")

    safe_title = re.sub(r'[^\w\-_ ]', '_', title)[:60].strip()
    base_filename = f"{safe_title}_{conv_id[:8]}"

    if format == "md":
        lines = [f"# {title}", ""]
        if timestamp:
            lines += [f"**Datum:** {timestamp}", ""]
        if conv_type in ("session", "quiz") and difficulty:
            lines += [f"**Niveau:** {difficulty}", ""]
        lines += ["---", ""]
        for msg in messages:
            role_label = "Student" if msg["role"] in ("user", "student") else "Tutor"
            lines += [f"**{role_label}:** {msg['content']}", ""]
        if feedback:
            lines += ["---", "## Feedback", "", feedback]
        content = "\n".join(lines)
        return StreamingResponse(
            iter([content.encode("utf-8")]),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{base_filename}.md"'},
        )

    elif format == "docx":
        from docx import Document as DocxDocument
        from docx.shared import Pt

        doc = DocxDocument()
        doc.add_heading(title, level=1)
        if timestamp:
            doc.add_paragraph(f"Datum: {timestamp}")
        if conv_type in ("session", "quiz") and difficulty:
            doc.add_paragraph(f"Niveau: {difficulty}")
        doc.add_paragraph("")
        for msg in messages:
            role_label = "Student" if msg["role"] in ("user", "student") else "Tutor"
            p = doc.add_paragraph()
            run_label = p.add_run(f"{role_label}: ")
            run_label.bold = True
            run_label.font.size = Pt(11)
            p.add_run(msg["content"])
        if feedback:
            doc.add_paragraph("")
            doc.add_heading("Feedback", level=2)
            doc.add_paragraph(feedback)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base_filename}.docx"'},
        )

    elif format == "pdf":
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, title.encode("latin-1", errors="replace").decode("latin-1"), ln=True)
        pdf.set_font("Helvetica", "", 10)
        if timestamp:
            pdf.cell(0, 6, f"Datum: {timestamp}", ln=True)
        if conv_type in ("session", "quiz") and difficulty:
            pdf.cell(0, 6, f"Niveau: {difficulty}", ln=True)
        pdf.ln(4)

        for msg in messages:
            role_label = "Student" if msg["role"] in ("user", "student") else "Tutor"
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, f"{role_label}:", ln=True)
            pdf.set_font("Helvetica", "", 10)
            safe_content = msg["content"].encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 5, safe_content)
            pdf.ln(2)

        if feedback:
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "Feedback", ln=True)
            pdf.set_font("Helvetica", "", 10)
            safe_fb = feedback.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 5, safe_fb)

        pdf_bytes = pdf.output()
        return StreamingResponse(
            iter([bytes(pdf_bytes)]),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{base_filename}.pdf"'},
        )


@app.delete("/api/conversations/{conv_id}")
async def delete_conversation(conv_id: str):
    """Delete a saved conversation."""
    for filepath in config.CONVERSATIONS_DIR.glob("*.json"):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            if data.get("id") == conv_id:
                os.remove(filepath)
                logger.info(f"Rapport verwijderd: {filepath.name}")
                return {"status": "ok"}
        except Exception:
            continue
    raise HTTPException(404, "Rapport niet gevonden.")


@app.post("/api/chat/save")
async def save_chat_history(req: ChatRequest):
    """Manually save the current chat history."""
    history = chat_histories.get(req.user_id, [])
    if not history:
        raise HTTPException(400, "Geen chatgeschiedenis om op te slaan.")
    # Use the first user message as topic
    first_msg = next((m["content"][:80] for m in history if m["role"] == "user"), "Chat")
    conv_id = _save_conversation("chat", req.user_id, history, {"topic": first_msg})
    # Clear the chat history after saving
    chat_histories[req.user_id] = []
    return {"status": "ok", "id": conv_id, "message_count": len(history)}


# ── Quiz endpoints ───────────────────────────────────────────────────


class QuizGenerateRequest(BaseModel):
    subject: str
    quiz_type: str  # "mc3", "mc4", "open"
    num_questions: int = 5
    difficulty: str = "basis"  # "basis", "gevorderd", "expert"


class QuizEvaluateRequest(BaseModel):
    questions: list
    answers: list
    quiz_type: str
    subject: str = ""


@app.post("/api/quiz/generate")
async def generate_quiz(req: QuizGenerateRequest):
    """Generate quiz questions using AI."""
    if not is_configured():
        raise HTTPException(503, "Geen AI API-key geconfigureerd. Ga naar Setup.")

    if req.num_questions < 1 or req.num_questions > 20:
        raise HTTPException(400, "Aantal vragen moet tussen 1 en 20 liggen.")

    quiz_type_label = {
        "mc3": "meerkeuzevragen met 3 antwoordmogelijkheden (A, B, C)",
        "mc4": "meerkeuzevragen met 4 antwoordmogelijkheden (A, B, C, D)",
        "open": "open vragen",
    }.get(req.quiz_type, "meerkeuzevragen met 4 antwoordmogelijkheden (A, B, C, D)")

    # Build RAG context if enabled
    rag_context = ""
    if rag_enabled:
        try:
            from rag.vector_store import search as rag_search
            context_items = rag_search(req.subject, n_results=10)
            if context_items:
                rag_context = "\n\n---\n\n".join(
                    f"[Bron: {item['source']}]\n{item['text']}" for item in context_items
                )
        except Exception as e:
            logger.warning(f"RAG context voor toets mislukt: {e}")

    if req.quiz_type in ("mc3", "mc4"):
        num_options = 3 if req.quiz_type == "mc3" else 4
        option_letters = "A, B, C" if num_options == 3 else "A, B, C, D"
        json_example = '{"questions": [{"question": "Vraag tekst hier?", "options": {"A": "optie A", "B": "optie B", "C": "optie C"}, "correct": "A"}]}'
        if num_options == 4:
            json_example = '{"questions": [{"question": "Vraag tekst hier?", "options": {"A": "optie A", "B": "optie B", "C": "optie C", "D": "optie D"}, "correct": "A"}]}'
        format_instruction = f"""Elke vraag heeft:
- "question": de vraagtekst
- "options": een object met {option_letters} als keys en de antwoordteksten als values
- "correct": de letter van het juiste antwoord ({option_letters})"""
    else:
        json_example = '{"questions": [{"question": "Vraag tekst hier?", "answer": "Het verwachte antwoord hier"}]}'
        format_instruction = """Elke vraag heeft:
- "question": de vraagtekst
- "answer": het verwachte modelantwoord (uitgebreid genoeg om te beoordelen)"""

    difficulty_instruction = {
        "basis": "Stel eenvoudige vragen die basisbegrippen en feitenkennis toetsen. Geschikt voor beginners.",
        "gevorderd": "Stel vragen van gevorderd niveau die zowel kennis als begrip, toepassing en analyse toetsen.",
        "expert": "Stel uitdagende vragen die diepgaand begrip, analyse en kritisch denken vereisen. Geschikt voor experts.",
    }.get(req.difficulty, "Stel eenvoudige vragen die basisbegrippen en feitenkennis toetsen.")

    system = f"""Je bent een toetsgenerator voor studenten. Genereer precies {req.num_questions} {quiz_type_label} over het onderwerp: "{req.subject}".

Moeilijkheidsgraad: {req.difficulty}.
{difficulty_instruction}

Antwoord UITSLUITEND met valid JSON in dit exacte formaat:
{json_example}

{format_instruction}

Regels:
- Genereer precies {req.num_questions} vragen.
- Stel duidelijke, relevante vragen die begrip toetsen.
- Alle vragen moeten op het niveau '{req.difficulty}' zijn.
- Bij meerkeuzevragen: maak de afleiders plausibel maar duidelijk onjuist.
- Antwoord ALLEEN met de JSON. Geen extra tekst, uitleg of markdown."""

    if rag_context:
        system += f"\n\nGebruik de volgende context uit de kennisbank als basis voor de vragen:\n\n{rag_context}"

    logger.info(f"Toets genereren: onderwerp='{req.subject}', type={req.quiz_type}, vragen={req.num_questions}, rag={rag_enabled}")

    try:
        result = create_chat_completion(
            messages=[{"role": "user", "content": f"Genereer een toets met {req.num_questions} {quiz_type_label} over: {req.subject}"}],
            system=system,
            return_usage=True,
        )
        raw_text = result["text"].strip()
        usage = result["usage"]

        # Try to extract JSON from the response
        # Sometimes the AI wraps it in markdown code blocks
        if raw_text.startswith("```"):
            # Remove markdown code block
            lines = raw_text.split("\n")
            json_lines = []
            in_block = False
            for line in lines:
                if line.strip().startswith("```") and not in_block:
                    in_block = True
                    continue
                elif line.strip().startswith("```") and in_block:
                    break
                elif in_block:
                    json_lines.append(line)
            raw_text = "\n".join(json_lines)

        quiz_data = json.loads(raw_text)
        questions = quiz_data.get("questions", [])

        logger.info(f"Toets gegenereerd: {len(questions)} vragen | tokens in={usage.get('input_tokens', '?')} out={usage.get('output_tokens', '?')}")

        return {
            "questions": questions,
            "quiz_type": req.quiz_type,
            "subject": req.subject,
            "usage": usage,
        }
    except json.JSONDecodeError as e:
        logger.error(f"Toets JSON parse fout: {e}\nRuwe tekst: {raw_text[:500]}")
        raise HTTPException(500, f"Kon de toetsvragen niet verwerken. Probeer opnieuw.")
    except RuntimeError as e:
        logger.error(f"Toets generatie fout: {e}")
        raise HTTPException(503, str(e))
    except Exception as e:
        logger.error(f"Onverwachte fout bij toets generatie: {e}")
        raise HTTPException(500, f"Er ging iets mis bij het genereren van de toets.")


@app.post("/api/quiz/evaluate")
async def evaluate_quiz(req: QuizEvaluateRequest):
    """Evaluate quiz answers using AI."""
    if not is_configured():
        raise HTTPException(503, "Geen AI API-key geconfigureerd. Ga naar Setup.")

    if not req.questions or not req.answers:
        raise HTTPException(400, "Vragen en antwoorden zijn vereist.")

    if req.quiz_type in ("mc3", "mc4"):
        # For MC questions, we can evaluate directly without AI
        results = []
        correct_count = 0
        for i, q in enumerate(req.questions):
            user_answer = req.answers[i] if i < len(req.answers) else ""
            is_correct = user_answer.upper() == q.get("correct", "").upper()
            if is_correct:
                correct_count += 1
            correct_letter = q.get("correct", "?")
            correct_text = q.get("options", {}).get(correct_letter, correct_letter)
            user_text = q.get("options", {}).get(user_answer.upper(), user_answer) if user_answer else "(niet beantwoord)"
            results.append({
                "question": q.get("question", ""),
                "user_answer": user_text,
                "user_answer_letter": user_answer.upper() if user_answer else "",
                "correct_answer": f"{correct_letter}: {correct_text}",
                "correct_letter": correct_letter,
                "is_correct": is_correct,
                "feedback": "Goed zo!" if is_correct else f"Het juiste antwoord is {correct_letter}: {correct_text}.",
            })

        # Get AI feedback on overall performance
        score_pct = round(correct_count / len(req.questions) * 100) if req.questions else 0
        try:
            wrong_questions = [r for r in results if not r["is_correct"]]
            feedback_prompt = f"De student maakte een meerkeuzetoets over '{req.subject}' en scoorde {correct_count}/{len(req.questions)} ({score_pct}%)."
            if wrong_questions:
                feedback_prompt += "\n\nFout beantwoorde vragen:\n"
                for r in wrong_questions:
                    feedback_prompt += f"- Vraag: {r['question']}\n  Antwoord student: {r['user_answer']}\n  Juiste antwoord: {r['correct_answer']}\n"
            feedback_prompt += "\nGeef per fout beantwoorde vraag een korte uitleg waarom het antwoord fout is en waarom het juiste antwoord correct is. Wees bemoedigend."

            fb_result = create_chat_completion(
                messages=[{"role": "user", "content": feedback_prompt}],
                system="Je bent een behulpzame tutor die feedback geeft op toetsresultaten. Communiceer in het Nederlands. Wees bemoedigend en constructief.",
                return_usage=True,
            )
            overall_feedback = fb_result["text"]
            usage = fb_result["usage"]
        except Exception as e:
            logger.warning(f"Toets feedback generatie mislukt: {e}")
            overall_feedback = ""
            usage = {"input_tokens": 0, "output_tokens": 0}

        return {
            "results": results,
            "score": correct_count,
            "total": len(req.questions),
            "percentage": score_pct,
            "overall_feedback": overall_feedback,
            "usage": usage,
        }
    else:
        # Open questions: use AI to evaluate
        eval_items = []
        for i, q in enumerate(req.questions):
            user_answer = req.answers[i] if i < len(req.answers) else "(niet beantwoord)"
            eval_items.append({
                "nr": i + 1,
                "question": q.get("question", ""),
                "model_answer": q.get("answer", ""),
                "user_answer": user_answer,
            })

        eval_json = json.dumps(eval_items, ensure_ascii=False)

        system = """Je bent een toetsbeoordelaar. Beoordeel de antwoorden van de student.

Antwoord UITSLUITEND met valid JSON in dit exacte formaat:
{"results": [{"nr": 1, "is_correct": true, "score": 1.0, "feedback": "uitleg"}], "overall_feedback": "samenvatting"}

Per vraag:
- "nr": het vraagnummer
- "is_correct": true als het antwoord (grotendeels) correct is, false als het fout is
- "score": 0.0 tot 1.0 (hoe goed het antwoord is)
- "feedback": korte feedback op het antwoord van de student

"overall_feedback": een samenvattend commentaar met tips.

Regels:
- Beoordeel op inhoud, niet op formulering.
- Een antwoord hoeft niet woordelijk overeen te komen met het modelantwoord.
- Wees eerlijk maar bemoedigend.
- Antwoord ALLEEN met JSON, geen extra tekst."""

        try:
            result = create_chat_completion(
                messages=[{"role": "user", "content": f"Beoordeel deze toetsantwoorden:\n\n{eval_json}"}],
                system=system,
                return_usage=True,
            )
            raw_text = result["text"].strip()
            usage = result["usage"]

            if raw_text.startswith("```"):
                lines = raw_text.split("\n")
                json_lines = []
                in_block = False
                for line in lines:
                    if line.strip().startswith("```") and not in_block:
                        in_block = True
                        continue
                    elif line.strip().startswith("```") and in_block:
                        break
                    elif in_block:
                        json_lines.append(line)
                raw_text = "\n".join(json_lines)

            eval_data = json.loads(raw_text)
            ai_results = eval_data.get("results", [])
            overall_feedback = eval_data.get("overall_feedback", "")

            # Build final results
            results = []
            correct_count = 0
            for i, q in enumerate(req.questions):
                ai_r = ai_results[i] if i < len(ai_results) else {}
                is_correct = ai_r.get("is_correct", False)
                if is_correct:
                    correct_count += 1
                user_answer = req.answers[i] if i < len(req.answers) else "(niet beantwoord)"
                results.append({
                    "question": q.get("question", ""),
                    "user_answer": user_answer,
                    "correct_answer": q.get("answer", ""),
                    "is_correct": is_correct,
                    "score": ai_r.get("score", 1.0 if is_correct else 0.0),
                    "feedback": ai_r.get("feedback", ""),
                })

            total = len(req.questions)
            score_pct = round(correct_count / total * 100) if total else 0

            return {
                "results": results,
                "score": correct_count,
                "total": total,
                "percentage": score_pct,
                "overall_feedback": overall_feedback,
                "usage": usage,
            }
        except json.JSONDecodeError as e:
            logger.error(f"Toets evaluatie JSON parse fout: {e}\nRuwe tekst: {raw_text[:500]}")
            raise HTTPException(500, "Kon de beoordeling niet verwerken. Probeer opnieuw.")
        except RuntimeError as e:
            logger.error(f"Toets evaluatie fout: {e}")
            raise HTTPException(503, str(e))
        except Exception as e:
            logger.error(f"Onverwachte fout bij toets evaluatie: {e}")
            raise HTTPException(500, "Er ging iets mis bij het beoordelen van de toets.")


@app.post("/api/quiz/save")
async def save_quiz_result(request: Request):
    """Save a quiz result as a conversation/rapport."""
    data = await request.json()
    subject = data.get("subject", "Toets")
    quiz_type = data.get("quiz_type", "mc4")
    difficulty = data.get("difficulty", "gemiddeld")
    score = data.get("score", 0)
    total = data.get("total", 0)
    percentage = data.get("percentage", 0)
    results = data.get("results", [])
    overall_feedback = data.get("overall_feedback", "")

    type_labels = {"mc3": "MC (3 opties)", "mc4": "MC (4 opties)", "open": "Open vragen"}
    type_label = type_labels.get(quiz_type, quiz_type)
    diff_labels = {"basis": "Basis", "gevorderd": "Gevorderd", "expert": "Expert"}
    diff_label = diff_labels.get(difficulty, difficulty)

    # Build messages array for the conversation
    messages = []
    messages.append({
        "role": "assistant",
        "content": f"**Toets: {subject}**\nType: {type_label} — Niveau: {diff_label}\nAantal vragen: {total}\n\n---",
    })

    for i, r in enumerate(results):
        # Question as tutor message
        q_text = f"**Vraag {i+1}:** {r.get('question', '')}"
        if "options" in r:
            for letter, text in r["options"].items():
                q_text += f"\n{letter}) {text}"
        messages.append({"role": "assistant", "content": q_text})

        # Student answer
        messages.append({"role": "user", "content": r.get("user_answer", "(niet beantwoord)")})

        # Feedback
        icon = "✅" if r.get("is_correct") else "❌"
        fb = f"{icon} "
        if not r.get("is_correct"):
            fb += f"Juiste antwoord: {r.get('correct_answer', '?')}\n"
        fb += r.get("feedback", "")
        messages.append({"role": "assistant", "content": fb})

    # Overall result
    messages.append({
        "role": "assistant",
        "content": f"**Resultaat: {score}/{total} correct ({percentage}%)**\n\n{overall_feedback}",
    })

    conv_id = _save_conversation("quiz", "web_user", messages, {
        "topic": f"Toets: {subject}",
        "quiz_subject": subject,
        "quiz_type": quiz_type,
        "quiz_difficulty": difficulty,
        "score": score,
        "total": total,
        "percentage": percentage,
    })

    return {"status": "ok", "id": conv_id}


# ── Web interface ────────────────────────────────────────────────────


@app.get("/api/health")
async def health_check():
    """Check if the AI provider is configured."""
    settings = load_settings()
    provider = settings.get("provider", "anthropic")
    if provider == "ollama":
        configured = True
    elif provider == "openrouter":
        configured = bool(settings.get("openrouter_api_key"))
    else:
        configured = bool(settings.get("ai_api_key"))
    return {
        "status": "ok" if configured else "no_key",
        "configured": configured,
        "provider": provider,
        "model": settings.get("model", config.CLAUDE_MODEL),
        "version": config.VERSION,
    }


class SettingsRequest(BaseModel):
    provider: str
    model: str
    ai_api_key: str = ""
    elevenlabs_api_key: str = ""
    ollama_base_url: str = ""
    openrouter_api_key: str = ""
    tavily_api_key: str = ""
    user_name: str = ""
    user_education: str = ""
    user_start_year: str = ""


@app.get("/api/settings")
async def get_settings():
    """Return current settings (keys are masked for security)."""
    s = load_settings()
    ai_key = s.get("ai_api_key", "")
    el_key = s.get("elevenlabs_api_key", "")
    masked_ai = (ai_key[:4] + "..." + ai_key[-4:]) if len(ai_key) > 8 else ("***" if ai_key else "")
    masked_el = (el_key[:4] + "..." + el_key[-4:]) if len(el_key) > 8 else ("***" if el_key else "")
    or_key = s.get("openrouter_api_key", "")
    masked_or = (or_key[:4] + "..." + or_key[-4:]) if len(or_key) > 8 else ("***" if or_key else "")
    provider = s.get("provider", "anthropic")
    if provider == "ollama":
        configured = True
    elif provider == "openrouter":
        configured = bool(or_key)
    else:
        configured = bool(ai_key)
    return {
        "provider": provider,
        "model": s.get("model", config.CLAUDE_MODEL),
        "configured": configured,
        "ai_api_key_set": bool(ai_key),
        "ai_api_key_masked": masked_ai,
        "elevenlabs_api_key_set": bool(el_key),
        "elevenlabs_api_key_masked": masked_el,
        "ollama_base_url": s.get("ollama_base_url", "http://localhost:11434"),
        "openrouter_api_key_set": bool(or_key),
        "openrouter_api_key_masked": masked_or,
        "tavily_api_key_set": bool(s.get("tavily_api_key", "")),
        "tavily_api_key_masked": (lambda k: (k[:4] + "..." + k[-4:]) if len(k) > 8 else ("***" if k else ""))(s.get("tavily_api_key", "")),
        "user_name": s.get("user_name", ""),
        "user_education": s.get("user_education", ""),
        "user_start_year": s.get("user_start_year", ""),
    }


@app.post("/api/settings")
async def update_settings(req: SettingsRequest):
    """Save provider, model, and API keys."""
    valid_providers = {"anthropic", "openai", "ollama", "openrouter"}
    if req.provider not in valid_providers:
        raise HTTPException(400, f"Ongeldige provider: {req.provider}")
    new_settings = {"provider": req.provider, "model": req.model}
    if req.ai_api_key:
        new_settings["ai_api_key"] = req.ai_api_key
    if req.elevenlabs_api_key:
        new_settings["elevenlabs_api_key"] = req.elevenlabs_api_key
    if req.ollama_base_url:
        new_settings["ollama_base_url"] = req.ollama_base_url
    if req.openrouter_api_key:
        new_settings["openrouter_api_key"] = req.openrouter_api_key
    if req.tavily_api_key:
        new_settings["tavily_api_key"] = req.tavily_api_key
    # User profile (always save, even if empty to allow clearing)
    new_settings["user_name"] = req.user_name.strip()
    new_settings["user_education"] = req.user_education.strip()
    new_settings["user_start_year"] = req.user_start_year.strip()
    save_settings(new_settings)
    logger.info(f"Instellingen bijgewerkt: provider={req.provider}, model={req.model}")
    return {"status": "ok"}


@app.get("/api/info")
async def get_app_info():
    """Return comprehensive application information."""
    import sys
    settings = load_settings()
    ai_key = settings.get("ai_api_key", "")
    el_key = settings.get("elevenlabs_api_key", "")

    # Count documents and chunks
    docs = list_documents()
    doc_files = [f for f in config.DOCUMENTS_DIR.iterdir() if f.is_file() and f.suffix.lower() in {
        ".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".htm"
    }] if config.DOCUMENTS_DIR.exists() else []

    # Count extracted images
    image_count = 0
    if config.IMAGES_DIR.exists():
        for d in config.IMAGES_DIR.iterdir():
            if d.is_dir():
                image_count += len(list(d.glob("*.png"))) + len(list(d.glob("*.jpg")))

    return {
        "app": {
            "name": "StudyPilot",
            "version": config.VERSION,
            "version_name": config.VERSION_NAME,
            "tagline": "Minder zoeken, meer simuleren en sneller leren!",
            "license": "MIT",
            "author": "twinology.ai",
            "repository": "https://github.com/twinology/studypilot",
        },
        "ai_provider": {
            "active_provider": settings.get("provider", "anthropic"),
            "active_model": settings.get("model", config.CLAUDE_MODEL),
            "ai_key_configured": bool(ai_key),
            "elevenlabs_key_configured": bool(el_key),
            "providers": {
                "anthropic": {
                    "name": "Anthropic",
                    "description": "Claude AI — state-of-the-art reasoning en analyse",
                    "models": ["Sonnet 4.6", "Opus 4.6", "Sonnet 4", "Opus 4", "Haiku 4", "3.5 Sonnet", "3.5 Haiku", "3 Opus"],
                    "features": ["Multimodal (vision)", "200K context window", "Tool use", "Streaming"],
                    "website": "https://console.anthropic.com",
                },
                "openai": {
                    "name": "OpenAI",
                    "description": "GPT-modellen — breed ecosysteem en integraties",
                    "models": ["GPT-5.2", "5.2 Pro", "GPT-5.1", "GPT-5", "5 Mini", "5 Nano", "GPT-4.1", "4.1 Mini", "GPT-4o", "o4 Mini", "o3", "o3 Pro", "o1", "Codex"],
                    "features": ["Responses API", "Multimodal (vision)", "128K+ context", "Reasoning (o-serie)"],
                    "website": "https://platform.openai.com",
                },
                "openrouter": {
                    "name": "OpenRouter",
                    "description": "Eén API-key voor 300+ modellen van alle providers",
                    "models": ["GPT-5.2", "Claude Sonnet 4.6", "Gemini 2.5", "Llama 4", "Grok", "Mistral", "en 300+ meer..."],
                    "features": ["300+ modellen", "Alle providers", "Eén API-key", "Pay-per-use"],
                    "website": "https://openrouter.ai",
                },
                "ollama": {
                    "name": "Ollama (Lokaal)",
                    "description": "Draai AI-modellen lokaal — privacy-first, geen API-key nodig",
                    "models": ["Llama 3.1", "Mistral", "Gemma 2", "Phi-3", "CodeLlama", "en meer..."],
                    "features": ["100% lokaal", "Geen API-key nodig", "Privacy-first", "Gratis"],
                    "website": "https://ollama.com",
                },
            },
        },
        "features": {
            "multimodal_rag": {
                "name": "Multimodal RAG",
                "description": "Retrieval-Augmented Generation met tekst én afbeeldingen uit documenten",
                "embedding_model": config.EMBEDDING_MODEL,
                "vector_db": "ChromaDB",
                "chunk_size": config.CHUNK_SIZE,
                "chunk_overlap": config.CHUNK_OVERLAP,
                "image_extraction": config.EXTRACT_IMAGES,
                "max_images_per_query": config.MAX_IMAGES_IN_CONTEXT,
            },
            "elevenlabs_tts": {
                "name": "ElevenLabs Text-to-Speech",
                "description": "Natuurlijke stemmen met emotie-modulatie voor realistische gespreksvoering",
                "model": config.ELEVENLABS_MODEL_ID,
                "voices": list(config.ELEVENLABS_VOICES.keys()),
                "emotions": list(config.ELEVENLABS_EMOTION_SETTINGS.keys()),
                "website": "https://elevenlabs.io",
            },
            "practice_sessions": {
                "name": "Gesprekssimulaties met AI-acteurs",
                "description": "Simuleer communicatievaardigheden met AI-personages die in rol blijven",
            },
            "conversation_export": {
                "name": "Rapport exporteren",
                "description": "Exporteer rapporten als Markdown, DOCX of PDF",
                "formats": ["Markdown (.md)", "Word (.docx)", "PDF (.pdf)"],
            },
        },
        "user_profile": {
            "name": settings.get("user_name", ""),
            "education": settings.get("user_education", ""),
            "start_year": settings.get("user_start_year", ""),
        },
        "stats": {
            "documents_uploaded": len(doc_files),
            "unique_documents_indexed": len(docs),
            "images_extracted": image_count,
        },
        "system": {
            "python_version": sys.version.split()[0],
            "framework": "FastAPI",
            "frontend": "Vanilla JS (single-page)",
        },
    }


@app.get("/api/ollama/models")
async def get_ollama_models():
    """Fetch available models from Ollama instance."""
    import httpx
    settings = load_settings()
    base_url = settings.get("ollama_base_url", "http://localhost:11434")
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"{base_url}/api/tags")
            resp.raise_for_status()
            data = resp.json()
            models = []
            for m in data.get("models", []):
                name = m.get("name", "")
                size_gb = round(m.get("size", 0) / (1024**3), 1)
                models.append({
                    "value": name,
                    "label": f"{name} ({size_gb}GB)" if size_gb else name,
                })
            return {"status": "ok", "models": models, "base_url": base_url}
    except Exception as e:
        return {"status": "error", "models": [], "error": str(e), "base_url": base_url}


@app.get("/api/openrouter/models")
async def get_openrouter_models():
    """Fetch available models from OpenRouter."""
    import httpx
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get("https://openrouter.ai/api/v1/models")
            resp.raise_for_status()
            data = resp.json()
            models = []
            for m in data.get("data", []):
                model_id = m.get("id", "")
                name = m.get("name", model_id)
                pricing = m.get("pricing", {})
                prompt_cost = float(pricing.get("prompt", 0)) * 1_000_000  # per 1M tokens
                models.append({
                    "value": model_id,
                    "label": f"{name}",
                    "cost_per_m": round(prompt_cost, 2),
                })
            return {"status": "ok", "models": models, "total": len(models)}
    except Exception as e:
        return {"status": "error", "models": [], "error": str(e)}


# ── Token Stats Persistence ─────────────────────────────────────────

def _load_token_stats() -> dict:
    """Load cumulative token stats from disk."""
    defaults = {"input": 0, "output": 0, "total": 0, "chunks": 0, "cost": 0, "calls": 0}
    if config.TOKEN_STATS_FILE.exists():
        try:
            with open(config.TOKEN_STATS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Ensure all keys exist
            for k, v in defaults.items():
                data.setdefault(k, v)
            return data
        except Exception as e:
            logger.warning(f"Kon token_stats.json niet laden: {e}")
    return defaults


def _save_token_stats(stats: dict) -> None:
    """Save cumulative token stats to disk."""
    try:
        with open(config.TOKEN_STATS_FILE, "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2)
    except Exception as e:
        logger.warning(f"Kon token_stats.json niet opslaan: {e}")


@app.get("/api/token-stats")
async def get_token_stats():
    """Return persisted cumulative token stats."""
    return _load_token_stats()


@app.post("/api/token-stats")
async def update_token_stats(request: Request):
    """Add session delta to cumulative token stats."""
    delta = await request.json()
    current = _load_token_stats()
    for key in ("input", "output", "total", "chunks", "cost", "calls"):
        current[key] = current.get(key, 0) + delta.get(key, 0)
    _save_token_stats(current)
    return current


# ── Exchange Rate ───────────────────────────────────────────────────

_exchange_rate_cache = {"rate": None, "timestamp": 0}


@app.get("/api/exchange-rate")
async def get_exchange_rate():
    """Fetch current USD to EUR exchange rate (cached for 1 hour)."""
    import time
    import httpx

    now = time.time()
    # Return cached rate if less than 1 hour old
    if _exchange_rate_cache["rate"] and (now - _exchange_rate_cache["timestamp"]) < 3600:
        return {"rate": _exchange_rate_cache["rate"], "cached": True}

    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            # Free exchange rate API
            resp = await client.get("https://api.exchangerate-api.com/v4/latest/USD")
            resp.raise_for_status()
            data = resp.json()
            rate = data.get("rates", {}).get("EUR", 0.92)
            _exchange_rate_cache["rate"] = rate
            _exchange_rate_cache["timestamp"] = now
            return {"rate": rate, "cached": False}
    except Exception as e:
        logger.warning(f"Wisselkoers ophalen mislukt: {e}")
        # Fallback rate
        fallback = _exchange_rate_cache["rate"] or 0.92
        return {"rate": fallback, "cached": True, "fallback": True}


@app.get("/")
async def serve_web():
    html_path = config.BASE_DIR / "web" / "index.html"
    content = html_path.read_text(encoding="utf-8")
    # Inject cache-busting meta tag + version comment to force reload
    import time as _time
    content = content.replace("<head>", f'<head>\n<meta http-equiv="Cache-Control" content="no-cache, no-store, must-revalidate">\n<!-- v{config.VERSION}-{_time.time()} -->', 1)
    return HTMLResponse(
        content=content,
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )



# Serve static files (images, etc.) from web/ directory
app.mount("/images", StaticFiles(directory=str(config.BASE_DIR / "web" / "images")), name="images")
# Serve extracted document images for multimodal RAG
app.mount("/api/images", StaticFiles(directory=str(config.IMAGES_DIR)), name="doc_images")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port)
